import os
import json
from io import BytesIO

from docx import Document
import nltk
from sentence_transformers import SentenceTransformer
from scipy.spatial.distance import cosine
from PIL import Image

nltk.download('punkt')

model = SentenceTransformer('BAAI/bge-m3')

TOKEN_LIMIT = 512
IMAGE_SIZE_THRESHOLD = (100, 100)


class MinIODocProcessor:
    def __init__(self,
                 minio_client,
                 bucket_name: str,
                 ):
        """
        Инициализация процессора документов с поддержкой MinIO

        Args:
            minio_endpoint: URL сервера MinIO
            access_key: Ключ доступа к MinIO
            secret_key: Секретный ключ к MinIO
            bucket_name: Название бакета
        """
        self.minio_client = minio_client
        self.bucket_name = bucket_name

    def count_tokens(self, text: str) -> int:
        """Подсчет токенов в тексте"""
        return len(text.split())

    def load_minio_file(self, object_path: str) -> bytes:
        """
        Загрузка файла из MinIO

        Args:
            object_path: Путь к объекту в MinIO

        Returns:
            Содержимое файла в виде байтов

        Raises:
            ValueError: Если файл не найден
        """
        try:
            result = self.minio_client.get_object(self.bucket_name, object_path)
            return result.data
        except Exception as e:
            raise ValueError(f"Ошибка при загрузке файла из MinIO: {str(e)}")

    def process_minio_document(self, object_path: str) -> str:
        """
        Обработка документа из MinIO

        Args:
            object_path: Путь к документу .docx в MinIO

        Returns:
            JSON строка с результатами обработки
        """
        try:
            # Загружаем файл из MinIO
            file_content = self.load_minio_file(object_path)

            # Создаем временный файл для обработки
            temp_file = BytesIO(file_content)

            # Извлекаем содержимое документа
            sentences, images, tables = self.extract_text_images_tables(temp_file)

            # Группируем предложения
            chunks = self.cluster_sentences(sentences)

            # Формируем результат
            result = []
            for chunk in chunks:
                result.append({
                    "text": chunk,
                    "images": images,
                    "tables": tables
                })

            return json.dumps(result, indent=4, ensure_ascii=False)

        except Exception as e:
            raise ValueError(f"Ошибка при обработке документа: {str(e)}")

    def extract_text_images_tables(self, doc_stream: BytesIO):
        """
        Извлечение текста, изображений и таблиц из документа

        Args:
            doc_stream: Поток данных документа

        Returns:
            Tuple содержащий список предложений, пути к изображениям и данные таблиц
        """
        doc = Document(doc_stream)
        sentences = []
        images = []
        tables = []

        # Извлечение текста
        for para in doc.paragraphs:
            sentences.extend(nltk.tokenize.sent_tokenize(para.text))

        # Извлечение изображений
        img_count = 0
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                img_part = doc.part.rels[rel].target_part
                img_data = img_part.blob

                # Сохраняем изображение во временную область памяти
                img_bytes = BytesIO(img_data)

                # Проверяем размер изображения
                with Image.open(img_bytes) as img:
                    if img.size[0] >= IMAGE_SIZE_THRESHOLD[0] and img.size[1] >= IMAGE_SIZE_THRESHOLD[1]:
                        # Сохраняем путь к изображению в памяти
                        images.append(f"memory_image_{img_count}")
                        img_count += 1

        # Извлечение таблиц
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return sentences, images, tables

    def cluster_sentences(self, sentences, threshold: float = 0.4):
        """
        Группировка предложений по смыслу

        Args:
            sentences: Список предложений
            threshold: Порог схожести для группировки

        Returns:
            Список сгруппированных предложений
        """
        chunks = []
        if not sentences:
            return chunks

        current_chunk = [sentences[0]]
        current_embedding = model.encode(sentences[0])
        current_token_count = self.count_tokens(sentences[0])

        for i in range(1, len(sentences)):
            sent_embedding = model.encode(sentences[i])
            similarity = 1 - cosine(current_embedding, sent_embedding)
            sent_token_count = self.count_tokens(sentences[i])

            if similarity >= threshold and (current_token_count + sent_token_count) <= TOKEN_LIMIT:
                current_chunk.append(sentences[i])
                current_embedding = model.encode(" ".join(current_chunk))
                current_token_count += sent_token_count
            else:
                chunks.append(" ".join(current_chunk))
                current_chunk = [sentences[i]]
                current_embedding = sent_embedding
                current_token_count = sent_token_count

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks